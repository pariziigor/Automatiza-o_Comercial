import os
import pdfplumber
import re
from docx import Document

def ler_documento_cliente(filepath):
    """Lê o arquivo do cliente (PDF ou DOCX) e retorna texto puro."""
    if not filepath:
        return ""
        
    ext = os.path.splitext(filepath)[1].lower()
    texto = ""
    
    try:
        if ext == '.pdf':
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    texto += (page.extract_text() or "") + "\n"
        elif ext == '.docx':
            doc = Document(filepath)
            # Texto dos parágrafos
            for para in doc.paragraphs:
                texto += para.text + "\n"
            # Texto das tabelas (importante para orçamentos)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        texto += cell.text + " "
                    texto += "\n"
    except Exception as e:
        raise Exception(f"Erro ao ler arquivo: {str(e)}")
            
    return texto

def extrair_placeholders_modelo(docx_path):
    #Lê o arquivo Word e retorna uma LISTA de placeholders na ordem exata em que aparecem.
    #Remove duplicatas mantendo a primeira ocorrência.
    doc = Document(docx_path)
    placeholders_ordenados = []
    seen = set()

    # Regex para capturar {{ variavel }}
    # Aceita espaços dentro: {{ NOME }} ou {{NOME}}
    regex = r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}"

    def processar_texto(texto):
        matches = re.findall(regex, texto)
        for match in matches:
            if match not in seen:
                seen.add(match)
                placeholders_ordenados.append(match)

    # 1. Varre parágrafos do corpo principal
    for para in doc.paragraphs:
        processar_texto(para.text)

    # 2. Varre todas as tabelas (linha por linha, célula por célula)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    processar_texto(para.text)
                    
    # 3. Varre Cabeçalhos e Rodapés (Opcional, mas recomendado)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    processar_texto(para.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    processar_texto(para.text)

    return placeholders_ordenados